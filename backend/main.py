import os
import json
import asyncio
import time
import io
import unicodedata
import redis.asyncio as redis
from fastapi import FastAPI, HTTPException, Response
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

load_dotenv()

from app.state import AgentState
from app.planner import planner_node
from app.search_agent import search_agent_node
from app.scraper import scrape_url_content
from app.rag_agent import rag_retriever_node
from app.writer_agent import writer_node
from app.critic_agent import critic_node
from app.editor_agent import editor_node

app = FastAPI(title="Multi-Agent Streaming Research Engine")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Redis Client
redis_client = redis.Redis(host='localhost', port=6379, db=0, decode_responses=True)

class ResearchRequest(BaseModel):
    topic: str

class ExportRequest(BaseModel):
    topic: str
    markdown: str

# --- UTILITY: CHARACTER CLEANER ---
def clean_text(text: str) -> str:
    """Removes non-ASCII artifacts that cause black squares in PDFs."""
    text = unicodedata.normalize('NFKD', text)
    text = text.encode('ascii', 'ignore').decode('ascii')
    return text

# --- ASYNC WORKER ---
async def execute_sub_question_pipeline(sub_question: str, q_idx: int, topic: str, log_queue: asyncio.Queue) -> tuple[str, str]:
    await log_queue.put({"type": "log", "message": f"🚀 [Worker {q_idx + 1}] Launching: {sub_question}"})
    
    local_state: AgentState = {
        "topic": topic, "plan": {"sub_questions": [sub_question]}, "current_sub_question_index": 0,
        "search_queries": [], "scraped_raw_data": [], "vector_store_status": "idle",
        "drafted_sections": {}, "critic_feedback": {}, "revision_count": 0, "final_report": ""
    }
    
    search_updates = await asyncio.to_thread(search_agent_node, local_state)
    sources = search_updates.get("scraped_raw_data", [])
    
    scraped_payloads = []
    for src in sources:
        # Run synchronous playwright in background thread
        full_text = await asyncio.to_thread(scrape_url_content, src["url"])
        scraped_payloads.append({"url": src["url"], "title": src["title"], "extracted_content_preview": full_text})
    local_state["scraped_raw_data"] = scraped_payloads
    
    rag_updates = await asyncio.to_thread(rag_retriever_node, local_state)
    local_state["current_context"] = rag_updates.get("current_context", [])
    
    approved = False
    # Max 1 revision for speed
    while not approved and local_state["revision_count"] < 1:
        writer_updates = await asyncio.to_thread(writer_node, local_state)
        local_state["drafted_sections"].update(writer_updates.get("drafted_sections", {}))
        
        critic_updates = await asyncio.to_thread(critic_node, local_state)
        local_state["critic_feedback"] = critic_updates.get("critic_feedback", {})
        
        approved = local_state["critic_feedback"].get("approved", False)
        if not approved:
            local_state["revision_count"] += 1
            await asyncio.sleep(1)
            
    final_draft = local_state["drafted_sections"].get(sub_question, "Failed to compile.")
    return sub_question, final_draft

async def bounded_pipeline_execution(sem: asyncio.Semaphore, q: str, idx: int, topic: str, log_queue: asyncio.Queue):
    async with sem:
        return await execute_sub_question_pipeline(q, idx, topic, log_queue)

# --- SSE STREAMING ENDPOINT ---
@app.post("/api/v1/generate-report")
async def generate_full_report(payload: ResearchRequest):
    async def event_stream():
        log_queue = asyncio.Queue()
        async def run_architecture():
            try:
                start_time = time.time()
                await log_queue.put({"type": "log", "message": f"========== STARTING: {payload.topic} =========="})
                
                cache_key = f"research_cache:{payload.topic.strip().lower()}"
                cached_report = None
                try: cached_report = await redis_client.get(cache_key)
                except: pass

                if cached_report:
                    await log_queue.put({"type": "log", "message": "⚡ [Memory] Cache hit! Bypassing agents..."})
                    await log_queue.put({"type": "done", "report_markdown": cached_report})
                    return 

                state: AgentState = {"topic": payload.topic, "plan": {}, "drafted_sections": {}, "revision_count": 0, "final_report": ""}
                
                planner_updates = await asyncio.to_thread(planner_node, state)
                state.update(planner_updates)
                sub_questions = state["plan"].get("sub_questions", [])
                
                sem = asyncio.Semaphore(2) 
                tasks = [bounded_pipeline_execution(sem, q, idx, payload.topic, log_queue) for idx, q in enumerate(sub_questions)]
                completed_tasks = await asyncio.gather(*tasks)

                for sub_question, section_draft in completed_tasks:
                    state["drafted_sections"][sub_question] = section_draft

                editor_updates = await asyncio.to_thread(editor_node, state)
                state.update(editor_updates)
                
                final_markdown = state["final_report"]
                try: await redis_client.setex(cache_key, 604800, final_markdown)
                except: pass 
                
                await log_queue.put({"type": "log", "message": f"🏁 ========== COMPLETE in {time.time() - start_time:.2f}s =========="})
                await log_queue.put({"type": "done", "report_markdown": final_markdown})
            except Exception as e:
                await log_queue.put({"type": "error", "message": str(e)})

        asyncio.create_task(run_architecture())
        while True:
            event = await log_queue.get()
            yield f"data: {json.dumps(event)}\n\n"
            if event["type"] in ["done", "error"]: break

    return StreamingResponse(event_stream(), media_type="text/event-stream")

# --- EXPORT ENDPOINTS ---
@app.post("/api/v1/export/pdf")
async def export_pdf(payload: ExportRequest):
    clean_markdown = clean_text(payload.markdown)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"Research Report: {payload.topic}", styles['Title']), Spacer(1, 12)]
    
    for line in clean_markdown.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): story.append(Paragraph(line[2:], styles['Heading1']))
        elif line.startswith('## '): story.append(Paragraph(line[3:], styles['Heading2']))
        else: story.append(Paragraph(line, styles['BodyText']))
        
    doc.build(story)
    return Response(content=buffer.getvalue(), media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="Report.pdf"'})

@app.post("/api/v1/export/word")
async def export_word(payload: ExportRequest):
    clean_markdown = clean_text(payload.markdown)
    doc = Document()
    doc.add_heading(f"Research Report: {payload.topic}", 0)
    for line in clean_markdown.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        else: doc.add_paragraph(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return Response(content=buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="Report.docx"'})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)